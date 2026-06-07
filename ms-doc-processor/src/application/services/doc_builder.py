"""
application/services/doc_builder.py — Construcción del documento Word del PDC.

build_document() es el único punto de entrada público. Recibe un dict con
los datos del plan (resultado de PlanData.model_dump()) y devuelve un
objeto Document de python-docx listo para serializar.

Las funciones privadas (_build_header, _build_area_table, etc.) implementan
el formato oficial del PDC boliviano: tabla de datos referenciales + tabla de
desarrollo por área con fusiones verticales y colores alternados.
"""

import ast
import json
import re
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ─── Constantes ──────────────────────────────────────────────────────────────

MESES_ES = {
    1: "ENERO", 2: "FEBRERO", 3: "MARZO", 4: "ABRIL",
    5: "MAYO", 6: "JUNIO", 7: "JULIO", 8: "AGOSTO",
    9: "SEPTIEMBRE", 10: "OCTUBRE", 11: "NOVIEMBRE", 12: "DICIEMBRE",
}

TRIMESTRE_LITERAL = {1: "PRIMERO", 2: "SEGUNDO", 3: "TERCERO"}

COLOR_SOM   = "E2EFD9"  # verde claro — filas de encabezado y semanas impares
COLOR_WHITE = "FFFFFF"


# ─── Utilidades XML de python-docx ───────────────────────────────────────────

def set_cell_background(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     color="000000", size="4") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, enabled in [("top", top), ("bottom", bottom),
                           ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if enabled:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_vertical_align(cell, align: str = "top") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def merge_cells_horizontal(row, start_col: int, end_col: int):
    a = row.cells[start_col]
    b = row.cells[end_col]
    a.merge(b)
    return a


def set_column_width(col_cells, width_cm: float) -> None:
    # 1 cm ≈ 567 twips (unidad interna de Word)
    for cell in col_cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width_cm * 567)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


def paragraph_in_cell(cell, text: str, bold=False, font_size=9,
                       font_name="Arial Narrow",
                       align=WD_ALIGN_PARAGRAPH.LEFT,
                       space_before=0, space_after=0,
                       color=None, italic=False, underline=False):
    if cell.paragraphs and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


def add_paragraph(doc, text: str, bold=False, font_size=11,
                  font_name="Arial Narrow",
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6,
                  color=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


# ─── Helpers de formato ──────────────────────────────────────────────────────

def format_date_es(d: date) -> str:
    return f"{d.day} DE {MESES_ES[d.month]}"


def _extraer_periodos(carga_horaria: str) -> str:
    """'9 clases de 40 minutos a la semana' → '9'"""
    m = re.search(r"(\d+)\s+clases?", carga_horaria, re.IGNORECASE)
    return m.group(1) if m else "?"


# ─── Construcción del encabezado ─────────────────────────────────────────────

def _build_header(doc: Document, plan: dict) -> None:
    add_paragraph(doc, "EDUCACIÓN PRIMARIA COMUNITARIA VOCACIONAL",
                  bold=True, font_size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    add_paragraph(doc,
                  f"PLAN DE DESARROLLO CURRICULAR Nº {plan['numero_plan']}",
                  bold=True, font_size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    add_paragraph(doc, "1.    DATOS REFERENCIALES",
                  bold=True, font_size=10, space_before=2, space_after=4)

    tbl = doc.add_table(rows=0, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    col_widths = [3.5, 5.5, 4.0, 5.0]

    def ref_row(label1, value1, label2, value2):
        row = tbl.add_row()
        for i, (text, bold) in enumerate(
            [(label1, True), (value1, True), (label2, True), (value2, True)]
        ):
            c = row.cells[i]
            paragraph_in_cell(c, text, bold=bold, font_size=9)
            set_cell_vertical_align(c, "center")
        for i, w in enumerate(col_widths):
            set_column_width([row.cells[i]], w)
        return row

    def ref_row_span(label, value):
        row = tbl.add_row()
        c0 = row.cells[0]
        paragraph_in_cell(c0, label, bold=True, font_size=9)
        set_cell_vertical_align(c0, "center")
        set_column_width([c0], col_widths[0])
        c1 = merge_cells_horizontal(row, 1, 3)
        paragraph_in_cell(c1, value.upper(), bold=True, font_size=9)
        set_cell_vertical_align(c1, "center")
        return row

    fecha_desde = format_date_es(plan["fecha_inicio"])
    fecha_hasta = format_date_es(plan["fecha_fin"])
    trimestre   = TRIMESTRE_LITERAL.get(plan["trimestre"], str(plan["trimestre"]))

    # Construir la cadena de áreas a partir del JSONB, igual que antes pero
    # sin la columna areas ya eliminada del schema.
    contenido_hdr = plan["contenido"]
    if isinstance(contenido_hdr, str):
        contenido_hdr = json.loads(contenido_hdr)
    areas_str = " / ".join(
        a.get("nombre", "") for a in contenido_hdr.get("areas", [])
    )

    # El nombre del maestro se arma desde los dos campos libres del plan.
    titulo  = plan.get("titulo_docente") or ""
    maestro_str = f"{titulo} {plan.get('nombre_docente') or ''}".strip()

    ref_row("Distrito educativo", plan.get("distrito") or "",
            "Unidad educativa",   plan.get("unidad_educativa") or "")
    ref_row("Nivel",              "",
            "Año de escolaridad", plan["anio_escolaridad"].upper())
    ref_row_span("Director/a", plan.get("nombre_director") or "")
    ref_row_span("Maestro/a",  maestro_str)
    ref_row_span("Áreas",      areas_str)

    row = tbl.add_row()
    c0 = row.cells[0]
    paragraph_in_cell(c0, "Trimestre", bold=True, font_size=9)
    set_column_width([c0], col_widths[0])
    c1 = merge_cells_horizontal(row, 1, 3)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(trimestre.upper())
    r1.bold = True; r1.font.name = "Arial Narrow"; r1.font.size = Pt(9)
    r2 = p.add_run(f"     Del: {fecha_desde}     al: {fecha_hasta}")
    r2.bold = True; r2.font.name = "Arial Narrow"; r2.font.size = Pt(9)

    doc.add_paragraph()


# ─── Construcción de la tabla de desarrollo por área ─────────────────────────

def _build_area_table(doc: Document, area: dict) -> None:
    """
    Estructura de filas:
      0        : Nombre del área (colspan=6, fondo som)
      1        : Encabezados de columna (fondo som)
      2…2+n-1  : Una fila por semana (fondo alternado blanco/som)
      2+n      : Adaptaciones curriculares (cols 1-4 fusionadas)

    Columnas:
      0  Objetivo de aprendizaje  (fusionado vertical filas 2…adapt)
      1  Contenidos               (semana + temas)
      2  Momentos del proceso formativo
      3  Recursos
      4  Períodos
      5  Criterios de evaluación  (fusionado vertical filas 2…adapt)
    """
    area_nombre  = area.get("nombre", "")
    objetivo     = area.get("objetivo_aprendizaje", "")
    semanas      = area.get("semanas", [])
    criterios    = area.get("criterios_evaluacion", {})
    carga        = area.get("carga_horaria", "")
    adaptaciones = area.get(
        "adaptaciones_curriculares",
        "No se tienen estudiantes con dificultades de aprendizaje ni otro caso especial.",
    )

    n_semanas  = len(semanas)
    total_rows = 2 + n_semanas + 1
    col_w = [3.2, 2.8, 5.5, 2.2, 1.0, 3.0]

    tbl = doc.add_table(rows=total_rows, cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Fila 0: nombre del área
    row0 = tbl.rows[0]
    c_area = merge_cells_horizontal(row0, 0, 5)
    set_cell_background(c_area, COLOR_SOM)
    set_cell_vertical_align(c_area, "center")
    p_area = c_area.paragraphs[0]
    p_area.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_area = p_area.add_run(f"Área de saberes y conocimientos:  {area_nombre}")
    r_area.bold = True; r_area.font.name = "Arial Narrow"; r_area.font.size = Pt(9)

    # Fila 1: encabezados de columna
    row1 = tbl.rows[1]
    for i, hdr in enumerate([
        "Objetivo de aprendizaje", "Contenidos",
        "Momentos del proceso formativo",
        "Recursos", "Períodos", "Criterios de evaluación",
    ]):
        c = row1.cells[i]
        set_cell_background(c, COLOR_SOM)
        set_cell_vertical_align(c, "center")
        paragraph_in_cell(c, hdr, bold=True, font_size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

    periodos = _extraer_periodos(carga)

    for idx, semana in enumerate(semanas):
        srow = tbl.rows[2 + idx]
        bg   = COLOR_SOM if idx % 2 == 1 else COLOR_WHITE

        num_sem    = semana.get("numero", idx + 1)
        tema_sem   = semana.get("tema", "")
        practica   = semana.get("practica", "")
        teoria     = semana.get("teoria", "")
        valoracion = semana.get("valoracion", "")
        produccion = semana.get("produccion", "")
        materiales = semana.get("materiales", [])

        # Col 0: objetivo (se llenará solo en la primera fila; luego se fusiona)
        c0 = srow.cells[0]
        set_cell_background(c0, COLOR_WHITE)
        set_cell_vertical_align(c0, "top")
        if idx == 0:
            paragraph_in_cell(c0, objetivo, font_size=8,
                              align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Col 1: semana + temas
        c1 = srow.cells[1]
        set_cell_background(c1, bg)
        set_cell_vertical_align(c1, "top")
        p_sem = c1.paragraphs[0]
        p_sem.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_sem.paragraph_format.space_after = Pt(2)
        r_sem = p_sem.add_run(f"Semana {num_sem}:")
        r_sem.bold = True; r_sem.font.name = "Arial Narrow"; r_sem.font.size = Pt(8)

        temas_lista = tema_sem
        if isinstance(tema_sem, str):
            try:
                temas_lista = ast.literal_eval(tema_sem)
            except (ValueError, SyntaxError):
                temas_lista = [tema_sem]
        if not isinstance(temas_lista, list):
            temas_lista = [temas_lista]
        for tema in temas_lista:
            p_t = c1.add_paragraph()
            p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_t = p_t.add_run(f"➢ {tema}.")
            r_t.font.name = "Arial Narrow"; r_t.font.size = Pt(8)

        # Col 2: momentos del proceso formativo
        c2 = srow.cells[2]
        set_cell_background(c2, bg)
        set_cell_vertical_align(c2, "top")
        for i_mom, (etiqueta, contenido_moment) in enumerate([
            ("PRÁCTICA", practica), ("TEORÍA", teoria),
            ("VALORACIÓN", valoracion), ("PRODUCCIÓN", produccion),
        ]):
            p_mom = c2.paragraphs[0] if i_mom == 0 else c2.add_paragraph()
            p_mom.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_mom.paragraph_format.space_before = Pt(0 if i_mom == 0 else 3)
            p_mom.paragraph_format.space_after  = Pt(1)
            r_etiq = p_mom.add_run(f"{etiqueta}: ")
            r_etiq.bold = True; r_etiq.font.name = "Arial Narrow"; r_etiq.font.size = Pt(8)
            r_cont = p_mom.add_run(contenido_moment)
            r_cont.font.name = "Arial Narrow"; r_cont.font.size = Pt(8)

        # Col 3: recursos / materiales
        c3 = srow.cells[3]
        set_cell_background(c3, bg)
        set_cell_vertical_align(c3, "top")
        for i_mat, mat in enumerate(materiales):
            p_mat = c3.paragraphs[0] if i_mat == 0 else c3.add_paragraph()
            p_mat.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_mat = p_mat.add_run(mat)
            r_mat.font.name = "Arial Narrow"; r_mat.font.size = Pt(8)

        # Col 4: períodos
        c4 = srow.cells[4]
        set_cell_background(c4, bg)
        set_cell_vertical_align(c4, "center")
        paragraph_in_cell(c4, periodos, font_size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # Col 5: criterios (se llenará solo en la primera fila; luego se fusiona)
        c5 = srow.cells[5]
        set_cell_background(c5, COLOR_WHITE)
        set_cell_vertical_align(c5, "top")
        if idx == 0:
            for i_dim, (dim, key) in enumerate([
                ("SER", "ser"), ("SABER", "saber"), ("HACER", "hacer"),
            ]):
                texto = criterios.get(key, "")
                p_dim = c5.paragraphs[0] if i_dim == 0 else c5.add_paragraph()
                p_dim.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_dim.paragraph_format.space_before = Pt(0 if i_dim == 0 else 3)
                p_dim.paragraph_format.space_after  = Pt(2)
                r_dim = p_dim.add_run(f"{dim}: ")
                r_dim.bold = True; r_dim.font.name = "Arial Narrow"; r_dim.font.size = Pt(8)
                r_txt = p_dim.add_run(texto)
                r_txt.font.name = "Arial Narrow"; r_txt.font.size = Pt(8)

    # Fila de adaptaciones curriculares
    adapt_idx = 2 + n_semanas
    adapt_row = tbl.rows[adapt_idx]
    set_cell_background(adapt_row.cells[0], COLOR_WHITE)
    set_cell_background(adapt_row.cells[5], COLOR_WHITE)
    c_adapt = merge_cells_horizontal(adapt_row, 1, 4)
    set_cell_vertical_align(c_adapt, "top")
    p_adapt = c_adapt.paragraphs[0]
    p_adapt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_lbl = p_adapt.add_run("ADAPTACIONES CURRICULARES: ")
    r_lbl.bold = True; r_lbl.font.name = "Arial Narrow"; r_lbl.font.size = Pt(8)
    r_txt = p_adapt.add_run(f"({adaptaciones}).")
    r_txt.font.name = "Arial Narrow"; r_txt.font.size = Pt(8)

    # Fusiones verticales: col 0 y col 5 desde la primera semana hasta adaptaciones
    if n_semanas >= 1:
        tbl.rows[2].cells[0].merge(tbl.rows[adapt_idx].cells[0])
        tbl.rows[2].cells[5].merge(tbl.rows[adapt_idx].cells[5])

    # Anchos de columna
    for row in tbl.rows:
        for i, w in enumerate(col_w):
            set_column_width([row.cells[i]], w)


# ─── Punto de entrada público ─────────────────────────────────────────────────

def build_document(plan: dict) -> Document:
    """
    Recibe un dict con los campos de PlanData y devuelve el Document Word.
    Llamar con: build_document(plan_data.model_dump())
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21.59)   # Letter 8.5"
    section.page_height   = Cm(27.94)   # Letter 11"
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    contenido = plan["contenido"]
    if isinstance(contenido, str):
        contenido = json.loads(contenido)
    areas = contenido.get("areas", [])

    _build_header(doc, plan)

    add_paragraph(doc, "2.    DESARROLLO",
                  bold=True, font_size=10, space_before=4, space_after=4)
    add_paragraph(doc, "Objetivo holístico de nivel",
                  bold=True, font_size=9, space_before=2, space_after=2)
    add_paragraph(doc, plan.get("objetivo_holistico") or "",
                  font_size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    for area in areas:
        _build_area_table(doc, area)
        doc.add_paragraph()

    return doc
